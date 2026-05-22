"""生成《下周前三项工作可行性具体方案》Word 文档。"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT = Path(__file__).resolve().parent.parent.parent / "下周工作可行性方案-平面图与地图MVP.docx"


def set_font(style, name="宋体", size=12):
    style.font.name = name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    doc.add_paragraph(text, style=style)


def main():
    doc = Document()
    set_font(doc.styles["Normal"])

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("校园导览系统 — 下周前三项工作\n可行性具体方案")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph("项目名称：校园导览与信息导航系统")
    doc.add_paragraph("对应阶段：第四周（承接第三周路网与 Dijkstra 成果）")
    doc.add_paragraph("编制说明：本方案仅覆盖下周计划前三项；第四项「路网补全校验」在数据录入时并行进行。")
    doc.add_paragraph()

    # --- 总览 ---
    heading(doc, "一、方案总览与可行性结论", 1)
    doc.add_paragraph(
        "第三周已完成：10 份分层分区路网 JSON（f1_a_graph.json … f5_b_graph.json）、"
        "Python indoor_nav 与 JS pathfind.js 双端 Dijkstra、index.json 与数据规范文档。"
        "第四周前三项目标是在此基础上建立「可看的平面地图 + 可点的导航 MVP」，"
        "不依赖自动定位，与《二维地图主导-分阶段推进方案》阶段 1 一致。"
    )
    p = doc.add_paragraph()
    p.add_run("总体可行性：高。").bold = True
    doc.add_paragraph(
        "理由：（1）技术栈与现有仓库一致（静态 HTML + 本地 server_main.py 提供 JSON/图片）；"
        "（2）算法已就绪，MVP 只需读图、标点、画折线；"
        "（3）1F A/B 节点量适中（约 14/25 个），适合首层试点。"
        "主要风险在数据质量（边权、连通）与坐标采集工时，可通过工具化点击取点降低。"
    )

    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    hdr = ["序号", "工作项", "可行性", "预估工时（组内）"]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
    rows = [
        ("1", "1F A/B 平面图 PNG + meta", "高", "0.5–1 人天/层区"),
        ("2", "map.html 单层 MVP", "高", "1.5–2 人天"),
        ("3", "关键节点 x_px/y_px", "高", "1–1.5 人天/层区（含取点工具）"),
    ]
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            table.rows[ri].cells[ci].text = val

    doc.add_paragraph()

    # --- 任务1 ---
    heading(doc, "二、任务 1：制作/导出各层平面图 PNG（优先 1F A/B）", 1)

    heading(doc, "2.1 目标", 2)
    bullet(doc, "为 1F A 座、1F B 座各提供一张固定像素尺寸的底图，供 map.html 显示。")
    bullet(doc, "在对应 f1_a_graph.json、f1_b_graph.json 的 meta 中记录图片路径与宽高，作为全图层坐标基准。")

    heading(doc, "2.2 推荐目录与命名", 2)
    bullet(doc, "目录：PanoramaProject/plans/（与 node_nav/data 并列，便于 server 静态访问）")
    bullet(doc, "文件：plans/f1_a.png、plans/f1_b.png（后续 f2_a.png … 同规则）")
    bullet(doc, "格式：PNG（无损，适合线条图）；若原图为 JPG，导出时统一转 PNG 并固定最终尺寸。")

    heading(doc, "2.3 meta 字段约定（写入 JSON）", 2)
    doc.add_paragraph("在 meta 中增加（与现有 building/floor/zone/units 并存）：")
    meta_table = doc.add_table(rows=5, cols=3)
    meta_table.style = "Table Grid"
    meta_table.rows[0].cells[0].text = "字段"
    meta_table.rows[0].cells[1].text = "示例"
    meta_table.rows[0].cells[2].text = "说明"
    meta_rows = [
        ("planImage", "plans/f1_a.png", "相对 PanoramaProject 根目录的路径"),
        ("planWidth", "1200", "图像宽度（像素），整数"),
        ("planHeight", "800", "图像高度（像素），整数"),
        ("coordSpace", "image_px", "固定取值 image_px，表示 x_px/y_px 相对该图左上角"),
    ]
    for i, row in enumerate(meta_rows, 1):
        for j, v in enumerate(row):
            meta_table.rows[i].cells[j].text = v

    heading(doc, "2.4 制作步骤（可执行）", 2)
    steps1 = [
        "获取底稿：学院提供的 CAD/PDF 平面图，或实地测量后手绘轮廓；无官方图时可用 PowerPoint/draw.io 按走廊走向描边。",
        "统一画布：选定宽度（建议 1000–1600px），高度按平面图比例计算，两区使用相同缩放比例尺便于日后对比。",
        "图层内容：墙体/走廊填充（浅色）、房间编号、门、楼梯、电梯、洗手间、连廊出口；不必标注全部 graph 节点。",
        "导出 PNG：关闭「按窗口缩放」；记录最终宽×高，填入 meta.planWidth / planHeight。",
        "验收：用看图软件查看属性尺寸与 meta 一致；在浏览器直接打开 http://localhost:8000/plans/f1_a.png 能加载。",
    ]
    for i, s in enumerate(steps1, 1):
        doc.add_paragraph(f"{i}. {s}")

    heading(doc, "2.5 注意事项", 2)
    for s in [
        "导出后禁止再改尺寸，否则已填 x_px/y_px 全部失效。",
        "A/B 两区若来自同一张总平面图，应裁成两张独立 PNG，各管各 meta，避免坐标系混用。",
        "图中北向建议标在角落，方便与「西北分叉」等节点 label 对照。",
    ]:
        bullet(doc, s)

    heading(doc, "2.6 可行性评估", 2)
    doc.add_paragraph("技术难度：低。无第三方 SDK，仅需图像编辑工具。阻塞因素：是否拿到可用平面图；若无，手绘简图仍可支撑 MVP 演示。")

    # --- 任务2 ---
    heading(doc, "三、任务 2：单层二维地图页 MVP（map.html）", 1)

    heading(doc, "3.1 目标", 2)
    bullet(doc, "用户打开 map.html 后选择楼层/分区（首版可写死 1F-A 或提供下拉）。")
    bullet(doc, "显示 plans/*.png 底图，叠加节点（有坐标者）、起终点、最短路径折线。")
    bullet(doc, "支持下拉或点击地图选起点/终点，调用 Dijkstra 后高亮路径并显示距离和途经节点名称。")

    heading(doc, "3.2 技术方案（与现有仓库对齐）", 2)
    bullet(doc, "部署：继续由 server_main.py 提供静态文件（与 panorama.html 同级），默认入口可改为 map.html（可选）。")
    bullet(doc, "数据：fetch node_nav/data/f1_a_graph.json；meta.planImage 决定底图 URL。")
    bullet(doc, "算路：首版推荐浏览器内嵌 pathfind.js（与 node_nav 同逻辑），避免跨域；二期可改为 Python /api/route。")
    bullet(doc, "渲染：容器 div + img 底图 + svg 覆盖层；circle 画点、polyline 画路径、line 画边（可选淡显）。")
    bullet(doc, "坐标换算：scaleX = imgClientWidth / meta.planWidth，屏幕坐标 = x_px * scaleX（Y 同理）。")

    heading(doc, "3.3 页面功能清单（MVP 范围）", 2)
    feat = doc.add_table(rows=1, cols=3)
    feat.style = "Table Grid"
    feat.rows[0].cells[0].text = "功能"
    feat.rows[0].cells[1].text = "优先级"
    feat.rows[0].cells[2].text = "说明"
    feat_rows = [
        ("加载 graph + 底图", "P0", "失败时提示 JSON/图片路径"),
        ("下拉选择起点/终点（按 label）", "P0", "选项来自 nodes[].label"),
        ("按钮「规划路径」", "P0", "调用 dijkstra，不可达时提示"),
        ("SVG 高亮路径折线", "P0", "仅连接有 x_px/y_px 的节点；缺坐标则列表显示路径文字"),
        ("点击地图设起点/终点", "P1", "命中半径约 12px 内最近节点"),
        ("楼层/区切换", "P1", "下拉 f1_a / f1_b，切换 JSON 与 PNG"),
        ("图例（起点/终点/路径）", "P2", "角标说明"),
    ]
    for row in feat_rows:
        cells = feat.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v

    heading(doc, "3.4 文件结构建议", 2)
    for line in [
        "PanoramaProject/map.html          — 主页面",
        "PanoramaProject/js/pathfind.js    — 自 node_nav 复制或软链",
        "PanoramaProject/js/map-view.js    — 底图缩放、画点、算路交互",
        "PanoramaProject/css/map.css       — 布局",
    ]:
        bullet(doc, line)

    heading(doc, "3.5 实现步骤（建议 3 天）", 2)
    for i, s in enumerate([
        "第 1 天：静态页加载 f1_a 底图 + JSON；下拉选点；控制台打印路径。",
        "第 2 天：接入 pathfind.js；SVG 绘制有坐标节点；polyline 画路径。",
        "第 3 天：点击取点交互；切换 f1_b；错误处理与简单样式。",
    ], 1):
        doc.add_paragraph(f"第 {i} 步：{s}")

    heading(doc, "3.6 验收标准", 2)
    for s in [
        "在 1F-A 上从「洗手间」到「A区大门」能算出路径并显示（与 python -m indoor_nav route 结果一致）。",
        "切换 1F-B 后加载另一张图与 JSON，能完成至少一条路径演示。",
        "窗口缩放后节点仍贴图（坐标换算正确）。",
    ]:
        bullet(doc, s)

    heading(doc, "3.7 风险与对策", 2)
    risk2 = doc.add_table(rows=1, cols=3)
    risk2.style = "Table Grid"
    risk2.rows[0].cells[0].text = "风险"
    risk2.rows[0].cells[1].text = "影响"
    risk2.rows[0].cells[2].text = "对策"
    for row in [
        ("路径上节点无坐标", "无法画折线", "文字列表显示 path_labels；优先给走廊节点标点"),
        ("JSON 不连通", "算路失败", "先用已校验的 f1_b；并行修 edges"),
        ("图片与坐标系不一致", "点漂移", "统一 planWidth/Height，禁止改图尺寸"),
    ]:
        c = risk2.add_row().cells
        for i, v in enumerate(row):
            c[i].text = v

    doc.add_paragraph()

    # --- 任务3 ---
    heading(doc, "四、任务 3：关键节点标注 x_px / y_px", 1)

    heading(doc, "4.1 目标", 2)
    bullet(doc, "在 nodes[] 中为需要在地图上显示的点增加 x_px、y_px（整数像素，相对 planImage 左上角）。")
    bullet(doc, "与 nodes[].id 严格一一对应；label 保持中文不变。")

    heading(doc, "4.2 哪些节点必须标点（1F 试点）", 2)
    must = doc.add_table(rows=1, cols=2)
    must.style = "Table Grid"
    must.rows[0].cells[0].text = "role / 类型"
    must.rows[0].cells[1].text = "说明"
    for row in [
        ("corridor", "分叉、尽头、连廊 link_to_a / link_to_b"),
        ("vertical", "楼梯、电梯"),
        ("facility", "洗手间、饮水机等 facilities 挂载点"),
        ("door", "至少标前门；后门若参与路径则标"),
    ]:
        c = must.add_row().cells
        c[0].text, c[1].text = row

    heading(doc, "4.3 推荐采集流程（第二步，依赖任务 1）", 2)
    for i, s in enumerate([
        "打开取点工具页 pick-coords.html：显示 plans/f1_a.png，点击地图输出 { id, x_px, y_px }。",
        "左侧列表勾选当前要标的 node id（或搜索 label），点击图上位置写入。",
        "导出为 JSON 片段或一键合并进 f1_a_graph.json（避免手抄出错）。",
        "两人交叉检查：任一边两端节点均有坐标时，路径折线不应穿墙（目视）。",
    ], 1):
        doc.add_paragraph(f"{i}. {s}")

    heading(doc, "4.4 节点 JSON 示例", 2)
    doc.add_paragraph(
        '{ "id": "washroom", "label": "洗手间", "floor": 1, "role": "facility", '
        '"x_px": 420, "y_px": 310 }'
    )

    heading(doc, "4.5 与任务 2 的衔接", 2)
    bullet(doc, "map.html 仅绘制含 x_px、y_px 的节点；算路仍用全图 nodes/edges。")
    bullet(doc, "画路径时：沿 path_node_ids 顺序，对有坐标的点连线；中间缺坐标点可跳过或用虚线提示「数据待补」。")

    heading(doc, "4.6 工时与分工建议", 2)
    doc.add_paragraph(
        "1F-A（约 14 节点）：若 70% 需标点，约 30–45 分钟（工具有序点击）。"
        "1F-B（约 25 节点）：约 1 小时。建议一人制图、一人标点、一人对照实地走线校验。"
    )

    heading(doc, "4.7 注意事项", 2)
    for s in [
        "标点位置在「门洞靠走廊一侧」，不是房间几何中心。",
        "不要用 edges 的 weight 反推坐标；显示与算路分离。",
        "提交前在 meta 中确认 planWidth/planHeight 与 PNG 一致。",
    ]:
        bullet(doc, s)

    doc.add_paragraph()

    # --- 依赖关系 ---
    heading(doc, "五、三项工作的依赖关系与推荐顺序", 1)
    doc.add_paragraph("任务 1 → 任务 3 → 任务 2 可并行部分开发：")
    bullet(doc, "必须先完成任务 1（有 PNG 与 meta 宽高），任务 3 才能开始标点。")
    bullet(doc, "任务 2 可在仅有 PNG、尚无坐标时启动（仅底图 + 下拉算路 + 文字路径）。")
    bullet(doc, "任务 3 完成一部分坐标后，任务 2 即可增量显示折线（建议每完成一区合并一次 JSON）。")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("推荐时间线（1 周）：").bold = True
    timeline = [
        "第 1–2 天：完成 f1_a、f1_b PNG 与 meta；搭建 pick-coords.html。",
        "第 3–4 天：标注 1F A/B 关键节点坐标；并行开发 map.html 骨架。",
        "第 5 天：联调算路 + 画线；录制演示路径；列出仍缺坐标的节点清单。",
    ]
    for t in timeline:
        bullet(doc, t)

    heading(doc, "六、交付物清单", 1)
    deliver = [
        "plans/f1_a.png、plans/f1_b.png",
        "更新后的 f1_a_graph.json、f1_b_graph.json（含 meta.plan* 与节点坐标）",
        "map.html + js/css 静态资源",
        "pick-coords.html（取点工具，可选但强烈建议）",
        "简短使用说明（如何启动 server、打开 map.html）",
    ]
    for d in deliver:
        bullet(doc, d)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    main()
