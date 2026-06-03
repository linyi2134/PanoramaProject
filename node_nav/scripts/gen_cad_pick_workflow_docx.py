#!/usr/bin/env python3
"""生成 CAD 取点协作流程简要说明（Word）。"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "CAD取点协作流程-简要说明.docx"


def set_doc_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def main() -> None:
    doc = Document()
    set_doc_font(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("B7 室内导航 · CAD 底图取点协作流程（简要）")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph("项目：知途 · PanoramaProject")
    doc.add_paragraph("仓库：https://github.com/linyi2134/PanoramaProject")
    doc.add_paragraph()

    add_heading(doc, "一、从 GitHub 获取项目", 2)
    add_bullet(doc, "安装 Git 后，在任意目录打开 PowerShell，执行：")
    doc.add_paragraph(
        "git clone https://github.com/linyi2134/PanoramaProject.git",
        style="No Spacing",
    )
    add_bullet(doc, "进入项目目录：cd PanoramaProject")
    add_bullet(doc, "以后更新代码：在项目目录执行 git pull")
    doc.add_paragraph()

    add_heading(doc, "二、运行本地服务", 2)
    add_bullet(doc, "需 Python 3.10+；建议在 PanoramaProject 目录下操作。")
    add_bullet(doc, "启动：python server_main.py")
    add_bullet(doc, "浏览器访问（勿用 file:// 直接打开 HTML）：")
    add_bullet(doc, "地图导航：http://localhost:8000/map.html")
    add_bullet(doc, "CAD 取点工具：http://localhost:8000/tools/pick-coords.html")
    doc.add_paragraph()

    add_heading(doc, "三、CAD 底图说明（PDF → PNG）", 2)
    add_bullet(doc, "根目录 PDF 已统一命名：1F-A.pdf … 5F-A.pdf（A 座）；1_2_5F-B.pdf、3_4F-B.pdf（B 座）。")
    add_bullet(doc, "仓库已含 plans/f*_a_cad.png、f*_b_cad.png；一般无需自己转换。")
    add_bullet(doc, "若 PDF 更新，在项目目录执行：python node_nav/scripts/export_all_cad_plans.py --dpi 400")
    doc.add_paragraph()

    add_heading(doc, "四、在取点页标注节点（组员操作）", 2)
    add_bullet(doc, "打开 pick-coords.html，选择分区（A 座 / B 座）和楼层（1F–5F）。")
    add_bullet(doc, "将每个彩色圆点拖到 CAD 图上对应走廊、门口或楼梯位置；可勾选「显示边」对照路网。")
    add_bullet(doc, "可用「上一个 / 下一个」或键盘 ← → 切换节点。")
    add_bullet(doc, "完成后点击「导出 cad_pick JSON」，得到例如 cad_pick_f1_a.json。")
    add_bullet(doc, "将 JSON 文件发给负责人，或放入项目的 map_data/ 文件夹。")
    doc.add_paragraph()

    add_heading(doc, "五、写回路网数据（负责人）", 2)
    add_bullet(doc, "把 cad_pick_*.json 放到 PanoramaProject/map_data/ 下。")
    add_bullet(doc, "在项目目录执行，例如：")
    doc.add_paragraph(
        "python node_nav/scripts/apply_cad_coords.py map_data/cad_pick_f1_a.json",
        style="No Spacing",
    )
    add_bullet(doc, "成功后，对应 node_nav/data/f*_graph.json 会更新 x_px、y_px 与底图尺寸。")
    doc.add_paragraph()

    add_heading(doc, "六、分工建议", 2)
    add_bullet(doc, "可按楼层分工：每人负责 1～2 层的 pick-coords 标注并导出 JSON。")
    add_bullet(doc, "文件名须与楼层、分区一致，如 cad_pick_f3_b.json、cad_pick_f2_a.json。")
    add_bullet(doc, "标注前先在 CAD/PDF 上确认房间与走廊位置；旧示意图坐标已作废，以 CAD 为准。")
    doc.add_paragraph()

    add_heading(doc, "七、常见问题", 2)
    add_bullet(doc, "取点页 404：地址应为 pick-coords.html（不是 .htm），且必须先运行 server_main.py。")
    add_bullet(doc, "底图加载失败：执行 git pull，确认 plans/ 下有对应 *_cad.png。")
    add_bullet(doc, "git push 显示 Everything up-to-date：说明没有新 commit，需先 git add 与 git commit 再 push。")

    doc.save(OUT)
    print(f"已生成 {OUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
