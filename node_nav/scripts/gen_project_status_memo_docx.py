"""生成《项目进展备忘：已完成与待办》Word 文档。"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

OUT = Path(__file__).resolve().parent.parent.parent / "项目进展备忘-已完成与待办.docx"


def set_doc_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    doc = Document()
    set_doc_font(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("B7 校园导览与信息导航系统\n项目进展备忘")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"（已完成 / 未完成）· 更新日期：{date.today().isoformat()}")

    doc.add_paragraph()
    doc.add_paragraph(
        "仓库：https://github.com/linyi2134/PanoramaProject  ·  "
        "运行目录：PanoramaProject（python server_main.py）"
    )
    doc.add_paragraph(
        "产品定位：二维地图为主入口，360° 全景为辅；不做自动室内定位（手动选点/搜索）。"
    )

    # ── 已完成 ──
    add_heading(doc, "一、已完成", 1)

    add_heading(doc, "1. 二维地图（主产品）", 2)
    add_bullets(
        doc,
        [
            "map.html：B7 主楼 B 区 1–5F 示意导航（点击选点、路径高亮、侧栏房间列表、「最近厕所」）。",
            "已合并组员 final_map：扩展为 13 个楼层标签（主楼 1–5F + 小楼 1–3F + A 区办公楼 1–5F）。",
            "跨楼/跨层连通：CROSS 支持主楼楼梯、B 区连廊接 A 区、主楼接小楼、A 区接小楼等示意路径。",
            "算路引擎：浏览器端 js/pathfind.browser.js（与 Python indoor_nav、node_nav 同 Dijkstra 逻辑）。",
            "示意底图：plans/f1_b.png … f5_b.png（760×520）；部分 CAD 裁切图 f1_cad.png、f3_cad.png 等。",
            "1F B 区 JSON 试点：map_data/id_map_f1_b.json + f1_b_graph.json（x_px/y_px、meta.planImage）；"
            "map.html 的 JSON_GRAPHS 仅启用 1F；跨楼边在 JSON 模式下映射 fork_se、link_to_a。",
        ],
    )

    add_heading(doc, "2. 全景漫游（辅助）", 2)
    add_bullets(
        doc,
        [
            "panorama_full.html：52 场景（一栋/二栋/三栋/连廊），Pannellum + 热点跳转 + BFS 导航面板。",
            "panoramas/：52 张全景 jpg 已纳入仓库（体积约 300MB+，clone 较慢）。",
            "panorama.html：5 场景 demo 保留；与 map.html、panorama_full.html 互链。",
            "tools/build_panorama_full.py：可从组员 panorama(1).html 重新生成 panorama_full.html。",
        ],
    )

    add_heading(doc, "3. 路网数据与工具（课程/工程）", 2)
    add_bullets(
        doc,
        [
            "node_nav/data/：f1_a … f5_b 共 10 份 graph.json + index.json + 数据规范 README。",
            "Python 包 indoor_nav：命令行 route / nearest（示例：1F B 区 room137_front → washroom）。",
            "脚本：apply_layout_coords.py、export_map_png.py、dwg_or_pdf_to_png.py；"
            "tools/export_floor_png.html 浏览器导出示意 PNG/SVG。",
            "map_data/README.md：示意图 FLOORS 与 JSON 双轨说明及 PNG 坐标流程。",
        ],
    )

    add_heading(doc, "4. 文档与协作", 2)
    add_bullets(
        doc,
        [
            "AGENT_HANDOFF.md、PROJECT_CONTEXT.md、README.md 已更新（入口、13 层、全景完整版）。",
            "本地 Git 提交：Integrate 13-floor map and 52-scene panorama（commit 3ec3422）。",
            "组员交付的 final_map .html、panorama(1).html、全景图/ 内容已并入 PanoramaProject，"
            "上级 indoor_navigator 根目录原件可删除（仅作备份）。",
        ],
    )

    add_heading(doc, "5. 运行方式（备忘）", 2)
    add_bullets(
        doc,
        [
            "cd PanoramaProject → python server_main.py → 默认打开 http://localhost:8000/map.html",
            "全景完整版：http://localhost:8000/panorama_full.html",
            "勿在 indoor_navigator 根目录启动服务；勿用 file:// 打开 HTML。",
        ],
    )

    # ── 未完成 ──
    add_heading(doc, "二、未完成 / 待办", 1)

    add_heading(doc, "1. 数据与二维地图（建议优先）", 2)
    add_bullets(
        doc,
        [
            "2F–5F B 区：建立 id_map_f2_b.json …，运行 apply_layout_coords，在 map.html JSON_GRAPHS 逐层启用。",
            "路网连通性校验：各层 edges 是否连通、isFinished 标记；python -m indoor_nav route 批量样例。",
            "FLOORS 与 JSON 单源：长期合并示意图 id（f1_wash）与 JSON id（washroom），或抽离 schematic JSON。",
            "A 区 / 小楼：map.html 有示意节点，但无对应 f*_graph.json，无法与 Python 课程数据统一校验。",
        ],
    )

    add_heading(doc, "2. 平面图与坐标", 2)
    add_bullets(
        doc,
        [
            "CAD 底图（f1_cad.png 等）与示意 760×520 坐标不通用，尚未切换为 CAD 主底图。",
            "pick-coords.html 等取点工具未做：在 CAD 图上点击输出 x_px/y_px 仍靠手工。",
            "plans/ 若仅有 SVG 无 PNG，需运行 export_map_png.py --all 生成各层 PNG。",
        ],
    )

    add_heading(doc, "3. 全景与一体化（可选）", 2)
    add_bullets(
        doc,
        [
            "二维房间与全景场景未绑点（如 138 微机室 ↔ 一栋-1f-x），两套导航图独立。",
            "全景 BFS 与地图 Dijkstra 不是同一套图数据，无法保证路径完全一致。",
        ],
    )

    add_heading(doc, "4. 工程与其它", 2)
    add_bullets(
        doc,
        [
            "GitHub push：若本机尚未 push 成功，需在 PanoramaProject 下执行 git push（含大体积 panoramas）。",
            "PyInstaller 打包：build_exe.py 已含 panorama_full，但 52 张图会使 exe 体积很大，需实测。",
            "根目录未跟踪的重复文件 f1_b.png 可删除（与 plans/f1_b.png 重复）。",
        ],
    )

    add_heading(doc, "三、明确不做（组内决议）", 2)
    add_bullets(
        doc,
        [
            "BLE / WiFi 自动室内定位（采用手动选点、搜索；可选二维码）。",
            "Qt / OpenGL 桌面客户端（采用 Web + Python 静态服务）。",
        ],
    )

    add_heading(doc, "四、建议下一步（备忘）", 2)
    add_bullets(
        doc,
        [
            "短期：校验 f1_a/f1_b 连通性 + 做 2F JSON 坐标（延续 1F 流程）。",
            "中期：2F–5F 逐层启用 JSON_GRAPHS；抽测跨楼路径是否缺 CROSS 边。",
            "按需：CAD 取点工具、二维↔全景绑点、Git LFS（若 panoramas 推送困难）。",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("说明：").bold = True
    p.add_run(
        "本备忘依据 PanoramaProject 仓库 2026-06 状态整理，"
        "细节以 AGENT_HANDOFF.md、PROJECT_CONTEXT.md 为准。"
    )

    doc.save(OUT)
    print(f"已生成：{OUT}")


if __name__ == "__main__":
    main()
