"""生成第三周进展报告 Word 文档。"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT = Path(__file__).resolve().parent.parent.parent / "每周进展报告-第3周.docx"


def set_doc_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def main() -> None:
    doc = Document()
    set_doc_font(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("【每周进展报告】")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph("周次：第三周")
    doc.add_paragraph("项目名称：校园导览与信息导航系统")
    doc.add_paragraph()

    add_heading(doc, "一、本周核心进展（3–5 条，可量化）", 2)
    items = [
        "完成楼内路网数据入库与规范化：整理 B7 教学楼 1–5 层、A/B 座共 10 份图数据（f1_a_graph.json … f5_b_graph.json），并生成 index.json 总览（各层节点/边数量、isFinished 状态）。",
        "实现最短路径算法（双语言）：Python 模块 indoor_nav（Dijkstra + 按设施类型找最近点 + 命令行）；Node/浏览器侧 node_nav/src/pathfind.js（同结构 JSON，逻辑一致）。",
        "统一数据规范并归档旧文件：制定 f{楼层}_{区}_graph.json 命名、description / isFinished / meta.building·zone 等字段约定，编写 node_nav/data/README.md 与 normalize_graphs.py 整理脚本。",
        "明确产品路线与定位策略：确定「二维地图 + 搜索起终点」为主入口、全景为辅；输出《二维地图主导-分阶段推进方案》《扩展方案与技术栈建议》；定位采用手动选点/搜索（可选二维码），暂不投入 BLE 等自动定位。",
        "代码托管更新：多次推送至 GitHub（linyi2134/PanoramaProject），组员可拉取最新路网与算法代码。",
    ]
    for i, t in enumerate(items, 1):
        doc.add_paragraph(f"{i}. {t}")

    add_heading(doc, "二、本周完成的成果", 2)
    outcomes = [
        ("☑", "代码/功能模块", [
            "indoor_nav/：Dijkstra、nearest_facility_by_type、python -m indoor_nav route|nearest",
            "node_nav/：pathfind.js、示例脚本 index.js / run-b1b.js",
            "本地服务 server_main.py + 全景 panorama.html（延续第二周，本周与远程合并保留）",
        ]),
        ("☐", "原型修改（如有）", [
            "二维地图页面（map.html）尚未开发；全景仍为可运行原型",
        ]),
        ("☑", "测试/验证", [
            "命令行对 f1_a_graph.json、f1_b_graph.json 等做过路径试算（如洗手间→大门、微机室前门→洗手间）",
        ]),
        ("☑", "文档/记录", [
            "node_nav/data/README.md、《二维地图主导-分阶段推进方案》、《扩展方案与技术栈建议》",
        ]),
        ("☑", "其他", [
            "路网数据 _archive/ 保留导入前文件名；讨论节点粒度、坐标标注两步走（先 PNG 后 xy）",
        ]),
    ]
    for mark, label, subs in outcomes:
        doc.add_paragraph(f"{mark} {label}")
        for s in subs:
            doc.add_paragraph(s, style="List Bullet")

    add_heading(doc, "三、原型-代码对照检查", 2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "原型功能"
    hdr[1].text = "实现状态"
    hdr[2].text = "是否调整"
    hdr[3].text = "调整说明"
    rows = [
        (
            "360° 全景虚拟漫游（多场景、热点跳转）",
            "☑ 已实现  ☐ 部分  ☐ 未实现",
            "☐ 是  ☑ 否",
            "使用 Pannellum + 本地 HTTP；与课程原 Qt/OpenGL 方案不同，属技术路线调整",
        ),
        (
            "楼内设施一键查找（Dijkstra + 路网 JSON）",
            "☐ 已实现  ☑ 部分  ☐ 未实现",
            "☑ 是  ☐ 否",
            "算法与 JSON 数据已有；尚未接入搜索 UI、二维地图展示",
        ),
        (
            "地图标注与个性化",
            "☐ 已实现  ☐ 部分  ☑ 未实现",
            "☐ 是  ☐ 否",
            "列为后续阶段",
        ),
        (
            "多层二维地图 + 起终点搜索 + 路径展示（新主流程）",
            "☐ 已实现  ☐ 部分  ☑ 未实现",
            "☑ 是  ☐ 否",
            "本周完成方案与数据基础；下阶段做 PNG 底图与 map.html",
        ),
        (
            "自动室内定位（GPS/BLE 等）",
            "☐ 已实现  ☐ 部分  ☑ 未实现",
            "☑ 是  ☐ 否",
            "调整为用户选点/搜索（可选二维码），降低维护成本",
        ),
    ]
    for r in rows:
        row = table.add_row().cells
        for i, val in enumerate(r):
            row[i].text = val

    add_heading(doc, "四、当前主要风险与卡点（最多 3 条）", 2)
    risks = [
        "路网数据未全部定稿：各层 isFinished 多为 false，连通性与边权仍需对照实地与平面图核对（尤其 1F-A 等边较少处）。",
        "主界面仍为全景入口：二维地图与搜索未开发，课程演示若要求「地图导航」需下周尽快出单层 MVP。",
        "节点平面坐标缺失：尚未在 JSON 中填写 x_px/y_px，无法在平面上画点与路径折线。",
    ]
    for i, t in enumerate(risks, 1):
        doc.add_paragraph(f"{i}. {t}")

    add_heading(doc, "五、下周明确计划（2–4 条）", 2)
    plans = [
        "制作/导出各层平面图 PNG（先 1F A/B），固定分辨率并在 meta 中记录 planImage、宽高。",
        "开发单层二维地图页 MVP（map.html）：底图 + 读 f*_graph.json + 点击/下拉选起终点 + 算路并高亮路径。",
        "为关键节点标注 x_px/y_px（门、楼梯、洗手间、连廊等），与 id 一一对应。",
        "继续补全/校验路网：按实地丈量更新 edges[].weight，将完成层标记 isFinished: true。",
    ]
    for i, t in enumerate(plans, 1):
        doc.add_paragraph(f"{i}. {t}")

    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run("填写说明：可根据组内分工在「一」中增删条目；提交前核对 GitHub 最新 commit。")
    nr.italic = True

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    main()
