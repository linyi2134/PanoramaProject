"""生成第五周进展报告 Word 文档（格式对齐第三周）。"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent.parent
OUT = ROOT / "每周进展报告-第5周.docx"
OUT_PARENT = ROOT.parent / "每周进展报告-第5周.docx"


def set_doc_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def main() -> None:
    doc = Document()
    set_doc_font(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("【每周进展报告】")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph("周次：第五周")
    doc.add_paragraph("项目名称：校园导览与信息导航系统（知途 · B7 教学楼）")
    doc.add_paragraph(
        "说明：第四周以平面图、CAD 源图与组员交付资料的数据搜集与整理为主，"
        "未形成可演示的增量功能；本周在既有路网与算法基础上完成主界面整合，系统已可端到端运行。"
    )
    doc.add_paragraph()

    add_heading(doc, "一、本周核心进展（3–5 条，可量化）", 2)
    items = [
        "完成二维导航主界面整合：将组员终稿 final_map 合并入 map.html，"
        "支持 B7 主楼 1–5F、小楼 1–3F、A 区办公楼 1–5F 共 13 个楼层切换，"
        "并实现跨楼连廊、楼梯的示意最短路径规划（Dijkstra，与 Python/JSON 同逻辑）。",
        "完成全景完整版接入：整合 52 个全景场景（一栋/二栋/三栋/连廊）至 panorama_full.html，"
        "图片资源纳入 panoramas/，支持热点漫游与 BFS 智能导航面板；与二维地图页互链。",
        "落地 1F B 区「PNG 底图 + JSON 路网」试点：通过 id_map_f1_b.json 与 apply_layout_coords.py "
        "写入 x_px/y_px，map.html 可在 1F 以课程数据算路（如 room137_front → washroom），"
        "并叠加 plans/f1_b.png 示意底图。",
        "完善工程化配套：新增/更新 pathfind.browser.js、export_map_png.py、dwg_or_pdf_to_png.py、"
        "export_floor_png.html、AGENT_HANDOFF.md、README 等；本地服务默认打开 map.html，"
        "一键启动即可演示二维 + 全景双入口。",
        "形成项目备忘与版本归档：编写《项目进展备忘-已完成与待办》Word 文档，"
        "并将主仓库提交至 GitHub（含 map、全景与 plans 等资源，组员可 clone 后直接运行）。",
    ]
    for i, t in enumerate(items, 1):
        doc.add_paragraph(f"{i}. {t}")

    add_heading(doc, "二、本周完成的成果", 2)
    outcomes = [
        ("☑", "代码/功能模块", [
            "map.html：13 层示意地图、起终点选点、路径高亮、侧栏房间列表、「最近厕所」、跨层/跨楼提示",
            "panorama_full.html + panoramas/（52 场景）；panorama.html 保留 5 场景 demo",
            "indoor_nav + js/pathfind.browser.js：浏览器与命令行双端 Dijkstra 可用",
            "server_main.py：http://localhost:8000 本地服务，默认 map.html 入口",
        ]),
        ("☑", "原型修改（如有）", [
            "二维地图由「计划中」升级为可演示主产品；全景由 demo 升级为完整版并与地图互通",
            "产品定位落实：二维为主、全景为辅；用户手动选点（不做 BLE 自动定位）",
        ]),
        ("☑", "测试/验证", [
            "本地启动 server_main.py 后，map.html / panorama_full.html 可正常浏览与算路",
            "1F B 区：python -m indoor_nav route f1_b_graph.json room137_front washroom 等样例通过",
            "抽测跨楼层（主楼楼梯）、跨楼（连廊至 A 区/小楼）示意路径可生成",
        ]),
        ("☑", "文档/记录", [
            "AGENT_HANDOFF.md、PROJECT_CONTEXT.md、map_data/README.md 更新",
            "项目进展备忘-已完成与待办.docx（组内备忘）",
        ]),
        ("☑", "其他", [
            "第四周搜集的 CAD/PDF、组员 HTML/全景图已并入 PanoramaProject 目录结构",
            "plans/ 各层示意 PNG/SVG 与部分 CAD 裁切图可用于答辩展示",
        ]),
    ]
    for mark, label, subs in outcomes:
        doc.add_paragraph(f"{mark} {label}")
        for s in subs:
            doc.add_paragraph(s, style="List Bullet")

    add_heading(doc, "三、原型-代码对照检查", 2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "原型功能"
    hdr[1].text = "实现状态"
    hdr[2].text = "是否调整"
    hdr[3].text = "调整说明"
    rows = [
        (
            "360° 全景虚拟漫游（多场景、热点跳转）",
            "☑ 已实现  ☐ 部分  ☐ 未实现",
            "☐ 是  ☑ 否",
            "完整版 52 场景已上线；技术路线为 Pannellum + Web，满足课程演示",
        ),
        (
            "楼内设施一键查找（Dijkstra + 路网 JSON）",
            "☑ 已实现  ☐ 部分  ☐ 未实现",
            "☐ 是  ☐ 否",
            "地图侧「最近厕所」+ 命令行 nearest；1F B 区已与 JSON 对齐，其余楼层持续优化",
        ),
        (
            "多层二维地图 + 起终点搜索 + 路径展示（主流程）",
            "☑ 已实现  ☐ 部分  ☐ 未实现",
            "☐ 是  ☐ 否",
            "13 层示意导航已可用；2F–5F 与 JSON 逐层对齐为收尾优化项",
        ),
        (
            "地图标注与个性化",
            "☐ 已实现  ☑ 部分  ☐ 未实现",
            "☐ 是  ☐ 否",
            "房间/设施节点与图例已标注；深度个性化（收藏、主题等）列为可选增强",
        ),
        (
            "自动室内定位（GPS/BLE 等）",
            "☐ 已实现  ☐ 部分  ☑ 未实现",
            "☑ 是  ☐ 否",
            "按组内决议改为手动选点/搜索，降低实现与维护成本，不影响主流程验收",
        ),
    ]
    for r in rows:
        row = table.add_row().cells
        for i, val in enumerate(r):
            row[i].text = val

    add_heading(doc, "四、当前主要风险与卡点（最多 3 条）", 2)
    risks = [
        "部分楼层仍为「示意图 + 内嵌节点」与「课程 JSON」双轨并行，"
        "目前仅 1F B 区完成 id 对照与像素坐标；2F–5F 对齐工作属优化而非阻塞项。",
        "A 区、小楼在地图可导航，但尚未建立与 node_nav/data 的 JSON 图一一对应，"
        "Python 侧无法对全区统一校验——若答辩强调「全楼同一数据源」，需后续补图或说明范围。",
        "全景资源体积较大（约 300MB+），仓库 clone 与首次加载略慢；"
        "可通过压缩图或按需下载进一步优化体验。",
    ]
    for i, t in enumerate(risks, 1):
        doc.add_paragraph(f"{i}. {t}")

    add_heading(doc, "五、下周明确计划（2–4 条）", 2)
    plans = [
        "延续 1F 试点：完成 2F B 区 id_map + apply_layout_coords，在 map.html 启用 JSON_GRAPHS[2]。",
        "数据质量收尾：抽查各层 graph 连通性，补充 python -m indoor_nav route 演示用例，"
        "将已校核楼层标记 isFinished: true。",
        "答辩演练与体验优化：整理启动说明（server_main → map / panorama_full），"
        "修复演示中发现的个别跨楼路径缺口；可选增加二维↔全景场景快捷入口。",
        "（可选）CAD 底图取点工具或 2F–5F 批量坐标导入，提升与真实平面图一致性。",
    ]
    for i, t in enumerate(plans, 1):
        doc.add_paragraph(f"{i}. {t}")

    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run(
        "填写说明：本周系统已具备可演示的完整闭环（二维导航 + 全景漫游 + 命令行算路）；"
        "下阶段工作以数据对齐与体验打磨为主。提交前可核对 GitHub：linyi2134/PanoramaProject。"
    )
    nr.italic = True

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    doc.save(OUT)
    print("wrote", OUT)
    try:
        import shutil
        shutil.copy2(OUT, OUT_PARENT)
        print("wrote", OUT_PARENT)
    except OSError as e:
        print("skip parent copy:", e)


if __name__ == "__main__":
    main()
